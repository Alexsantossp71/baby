from docx import Document
import os

def inspect_template(path):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return
    
    doc = Document(path)
    print(f"--- Template Inspection: {os.path.basename(path)} ---")
    
    # Check paragraphs for placeholders
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"P{i:03}: {text[:150]}")
            
    # Check tables (common in templates for group info)
    for i, table in enumerate(doc.tables):
        print(f"\n--- TABLE {i} ---")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"Row: {row_text}")

if __name__ == "__main__":
    template_path = r'f:\projetos_opencode\PI2016Django\base_teorica\Modelo_-_Relatorio_Parcial_oficial.docx'
    inspect_template(template_path)
