import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_relatorio_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Erro: {md_path} não encontrado.")
        return

    doc = Document()

    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
        else:
            # Processamento simples de negrito dentro da linha
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True

    doc.save(docx_path)
    print(f"Documento salvo em: {docx_path}")

if __name__ == "__main__":
    base_path = r'f:\projetos_opencode\PI2016Django\base_teorica'
    md_file = os.path.join(base_path, 'relatorio_parcial_pi.md')
    docx_file = os.path.join(base_path, 'Relatorio_Parcial_Permutas_Baby.docx')
    create_relatorio_docx(md_file, docx_file)
