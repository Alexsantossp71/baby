import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def add_cover_page(doc, title, group_id, members, orientador):
    # UNIVESP Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO – UNIVESP")
    run.bold = True
    run.font.size = Pt(14)

    for _ in range(5): doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(16)

    for _ in range(3): doc.add_paragraph()

    # Relatório Parcial Text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO PARCIAL")
    run.bold = True
    run.font.size = Pt(12)

    for _ in range(3): doc.add_paragraph()

    # Group and Members
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Grupo: {group_id}")
    run.bold = True
    
    for member in members:
        doc.add_paragraph(member)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Orientador: {orientador}")
    run.bold = True

    # Fill space to bottom
    for _ in range(10): doc.add_paragraph()

    # City/Year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026")
    doc.add_page_break()

def create_professional_relatorio(md_path, docx_path):
    doc = Document()
    set_margins(doc)
    
    # Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Content extraction for cover
    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    group_id = "PJI110 - A2026S1N1 - Grupo 4"
    members = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    orientador = "Fernando Pinto Martinez"

    add_cover_page(doc, title, group_id, members, orientador)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Simple MD to DOCX converter (skipping cover parts already handled)
    lines = content.split('\n')
    skip_mode = False
    
    for line in lines:
        line = line.strip()
        if line.startswith('# Relatório Parcial') or line.startswith('**Curso:**') or line.startswith('**Identificação') or line.startswith('**Polo:') or line.startswith('**Orientador:') or line.startswith('## Integrantes'):
            continue
        if line == '---':
            continue
        if line.startswith('- ') and '(' in line and 'RA:' in line: # Skip members list in body if already in cover
            continue

        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif line:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
            # Bold processing
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True

    doc.save(docx_path)
    print(f"Relatório Profissional salvo em: {docx_path}")

if __name__ == "__main__":
    md_file = r'f:\projetos_opencode\PI2016Django\base_teorica\relatorio_parcial_pi.md'
    docx_file = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V2.docx'
    create_professional_relatorio(md_file, docx_file)
