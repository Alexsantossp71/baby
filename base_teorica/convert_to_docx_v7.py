import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def add_spacer(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

def add_capa_v7(doc, title, members, locations):
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO – UNIVESP")
    run.bold = True
    run.font.size = Pt(12)

    add_spacer(doc, 2)

    # Authors
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.size = Pt(12)

    add_spacer(doc, 10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    add_spacer(doc, 14)

    # Bottom
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(locations + "\n2026")
    doc.add_page_break()

def add_folha_rosto_v7(doc, title, members, orientador, locations):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO")
    run.bold = True
    
    add_spacer(doc, 2)
    
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)

    add_spacer(doc, 8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    add_spacer(doc, 4)

    # Natureza del trabajo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    run = p.add_run("Relatório Parcial apresentado na disciplina de Projeto Integrador em Computação I, para o curso de Tecnologia em Processos Gerenciais/Ciência de Dados da Universidade Virtual do Estado de SÃO PAULO (UNIVESP).")
    run.font.size = Pt(10)
    
    add_spacer(doc, 1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    run = p.add_run(f"Orientador: {orientador}")
    run.font.size = Pt(10)
    run.bold = True

    add_spacer(doc, 10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(locations + "\n2026")
    doc.add_page_break()

def add_resumo_v7(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RESUMO")
    run.bold = True
    
    add_spacer(doc, 1)
    
    resumo_text = (
        "Este trabalho apresenta o desenvolvimento da plataforma 'Permutas Baby', uma solução de economia circular "
        "focada na troca de artigos infantis. A proposta utiliza o framework Django para criar um ambiente "
        "digital seguro onde famílias podem cadastrar produtos, realizar propostas de troca por escambo e gerenciar "
        "seu inventário. O projeto aplica metodologias de Design Thinking e Design Centrado no Humano para mitigar "
        "o desperdício de itens de alto giro como fraldas e roupas, promovendo sustentabilidade financeira e ambiental."
    )
    p = doc.add_paragraph(resumo_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    add_spacer(doc, 1)
    p = doc.add_paragraph()
    run = p.add_run("PALAVRAS-CHAVE: ")
    run.bold = True
    p.add_run("Economia Circular, Django, Permuta Colaborativa, Sustentabilidade.")
    doc.add_page_break()

def add_sumario_v7(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUMÁRIO")
    run.bold = True
    
    add_spacer(doc, 1)
    
    items = [
        ("1 INTRODUÇÃO", "05"),
        ("2 DESENVOLVIMENTO ITERATIVO", "06"),
        ("2.1 EMPATIA E DEFINIÇÃO", "06"),
        ("2.2 IDEAÇÃO E FUNCIONALIDADES", "06"),
        ("2.3 PROTOTIPAÇÃO E ARQUITETURA", "06"),
        ("3 FUNDAMENTAÇÃO TEÓRICA", "07"),
        ("4 CONSIDERAÇÕES PARCIAIS", "08"),
        ("5 REFERÊNCIAS BIBLIOGRÁFICAS", "09")
    ]
    
    for label, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        # Use tabs for beautiful dotted lines
        # In docx, we'd need to set tab stops. For simplicity and "beauty" we simulate with dots
        dots = "." * (80 - len(label))
        p.add_run(label)
        p.add_run(dots)
        p.add_run(page)
    
    doc.add_page_break()

def create_relatorio_v7(md_path, docx_path):
    doc = Document()
    set_margins(doc)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    members = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    orientador = "Fernando Pinto Martinez"
    locations = "Bragança Paulista, Campinas, Hortolândia, Jaguariúna, Mogi Mirim, Santa Bárbara D'Oeste & Santa Cruz das Palmeiras - SP"

    add_capa_v7(doc, title, members, locations)
    add_folha_rosto_v7(doc, title, members, orientador, locations)
    add_resumo_v7(doc, title)
    add_sumario_v7(doc)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    is_references = False
    skip_mode = False
    
    # Section counters
    main_section = 1
    
    for line in lines:
        line = line.strip()
        if not line: continue
            
        if line.startswith('# Relatório Parcial') or line.startswith('**Curso:') or line.startswith('**Identificação') or line.startswith('**Polo:') or line.startswith('**Orientador:'):
            continue
        
        if line.startswith('## Integrantes'):
            skip_mode = True
            continue
        if skip_mode:
            if line.startswith('## ') or line == '---':
                skip_mode = False
                if line == '---': continue
            else:
                continue
        
        if line == '---': continue

        if line.startswith('## '): # Major Section
            if "Referências" in line: is_references = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            
            # Simple manual numbering for major sections
            text = line[2:].strip()
            # Special case for References which doesn't get a number in some ABNT styles, but example has it as last.
            numbered_text = f"{main_section} {text.upper()}"
            if "REFERÊNCIAS" in text.upper(): numbered_text = text.upper()
            
            run = p.add_run(numbered_text)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            main_section += 1
            sub_section = 1
            
        elif line.startswith('### '): # Sub Section
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            text = line[3:].strip()
            run = p.add_run(f"{main_section-1}.{sub_section} {text}")
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sub_section += 1

        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
        else:
            p = doc.add_paragraph()
            if is_references:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(6)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Cm(1.25)
            
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True

    doc.save(docx_path)
    print(f"Relatório ABNT V7 (Final-Polish) salvo em: {docx_path}")

if __name__ == "__main__":
    md_file = r'f:\projetos_opencode\PI2016Django\base_teorica\relatorio_parcial_pi.md'
    docx_file = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V7_ABNT_Final.docx'
    create_relatorio_v7(md_file, docx_file)
