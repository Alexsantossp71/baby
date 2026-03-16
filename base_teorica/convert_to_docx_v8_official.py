import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def migrate_to_official_template(template_path, md_path, output_path):
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = Document(template_path)
    
    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    members_list = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    members_str = "\n".join(members_list)
    orientador = "Fernando Pinto Martinez"

    # --- COVER PAGE & FOLHA DE ROSTO REPLACEMENT ---
    for p in doc.paragraphs:
        # Title placeholder
        if "(Fonte: Arial ou Times 14)" in p.text:
            p.text = title.upper()
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Members placeholder
        if "Nome dos integrantes" in p.text:
            p.text = members_str
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Title placeholder 2 (Folha de Rosto)
        if "TÍTULO DO PROJETO" in p.text:
            p.text = title.upper()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Year
        if "202X" in p.text:
            p.text = p.text.replace("202X", "2026")
            
    # --- BODY CONTENT INJECTION ---
    # Identify indices of section headers first
    sections_found = {}
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip().upper()
        # Look for the UNIVESP Heading Style or standard numbering
        is_heading = p.style.name.startswith("Heading") or "TÍTULO NIVEL 1" in p.style.name.upper()
        if is_heading:
            if "INTRODUÇÃO" in txt: sections_found["INTRO"] = i
            elif "DESENVOLVIMENTO" in txt: sections_found["DEV"] = i
            elif "REFERÊNCIAS" in txt: sections_found["REF"] = i

    print(f"Sections identified at indices: {sections_found}")

    # Parse markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    md_sections = {}
    current_sec = None
    for line in md_text.split('\n'):
        line = line.strip()
        if line.startswith('## 2. Introdução'): current_sec = "INTRO"
        elif line.startswith('## 3. Desenvolvimento'): current_sec = "DEV"
        elif line.startswith('## 6. Referências'): current_sec = "REF"
        elif line.startswith('## '): current_sec = None # reset for other sections
        
        if current_sec and line and not line.startswith('## '):
            if current_sec not in md_sections: md_sections[current_sec] = []
            md_sections[current_sec].append(line)

    # Insert content
    # Note: We iterate backwards to avoid index shifting when deleting/inserting
    # However, for simplicity, we'll just find the paragraphs and replace/append.
    
    def clear_and_fill(idx, lines):
        # Find the header
        header_p = doc.paragraphs[idx]
        print(f"Feeding section starting after: {header_p.text}")
        
        # We'll just replace the first paragraph after the header if it looks like a placeholder
        # and append the rest. This is safer for "formatting".
        target_idx = idx + 1
        # Skip empty paragraphs after header
        while target_idx < len(doc.paragraphs) and not doc.paragraphs[target_idx].text.strip():
            target_idx += 1
            
        # Clear the instruction paragraph
        if target_idx < len(doc.paragraphs):
            doc.paragraphs[target_idx].text = ""
            
        for line in lines:
            new_p = doc.add_paragraph()
            # Try to copy style from a known "text-base" paragraph
            # We'll just use 'Normal' or 'a) texto-base' if it exists
            try:
                new_p.style = doc.styles['a) texto-base']
            except:
                new_p.style = doc.styles['Normal']
            
            new_p.text = line.replace('**', '') # Simple cleanup
            # Move it to the right position (after header)
            header_p._element.addnext(new_p._element)
            header_p = new_p # next line goes after this one

    for sec_key in ["INTRO", "DEV", "REF"]:
        if sec_key in sections_found and sec_key in md_sections:
            clear_and_fill(sections_found[sec_key], md_sections[sec_key])

    doc.save(output_path)
    print(f"Relatório Oficial V8 (Final) salvo em: {output_path}")

if __name__ == "__main__":
    t_path = r'f:\projetos_opencode\PI2016Django\base_teorica\Modelo_-_Relatorio_Parcial_oficial.docx'
    m_path = r'f:\projetos_opencode\PI2016Django\base_teorica\relatorio_parcial_pi.md'
    o_path = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V8_Oficial.docx'
    migrate_to_official_template(t_path, m_path, o_path)
