import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def add_cover_page_v5(doc, title, group_id, members, orientador):
    # 1. UNIVESP Header (Top)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO – UNIVESP")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph() # Small spacer

    # 2. Members (TOP - following example)
    for member in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Typical UNIVESP
        run = p.add_run(member)
        run.font.size = Pt(12)
    
    for _ in range(8): doc.add_paragraph() # Spacer to center

    # 3. Title (Center)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    for _ in range(3): doc.add_paragraph()

    # 4. Context info (Centered below title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO PARCIAL")
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Grupo: {group_id}")
    run.font.size = Pt(11)

    for _ in range(3): doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Orientador: {orientador}")
    run.bold = True

    # 5. Bottom (City/Year)
    for _ in range(10): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SÃO PAULO\n2026")
    doc.add_page_break()

def create_abnt_relatorio_v5(md_path, docx_path):
    doc = Document()
    set_margins(doc)
    
    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    group_id = "PJI110 - A2026S1N1 - Grupo 4"
    members = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    orientador = "Fernando Pinto Martinez"

    add_cover_page_v5(doc, title, group_id, members, orientador)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    is_references = False
    skip_mode = False
    
    for line in lines:
        line = line.strip()
        if not line: continue
            
        # --- ROBUST FILTERING ---
        # 1. Skip standard top headers
        if line.startswith('# Relatório Parcial') or line.startswith('**Curso:') or line.startswith('**Identificação') or line.startswith('**Polo:') or line.startswith('**Orientador:'):
            continue
        
        # 2. Skip the entire Integrantes section (Headings + list)
        if line.startswith('## Integrantes'):
            skip_mode = True
            continue
        
        # If we are in a skip section, keep skipping until we hit a horizontal rule or next section
        if skip_mode:
            if line.startswith('## ') or line == '---':
                skip_mode = False
                if line == '---': continue # Just skip the separator too
            else:
                continue
        
        if line == '---': continue # General separator skip
        # --- END FILTERING ---

        if line.startswith('## '):
            if "Referências" in line: is_references = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)

        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
        else:
            p = doc.add_paragraph()
            if is_references:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Cm(1.25)
            
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True

    doc.save(docx_path)
    print(f"Relatório ABNT V5 (Fixed Layout & Duplicate) salvo em: {docx_path}")

if __name__ == "__main__":
    md_file = r'f:\projetos_opencode\PI2016Django\base_teorica\relatorio_parcial_pi.md'
    docx_file = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V5_ABNT_Final.docx'
    create_abnt_relatorio_v5(md_file, docx_file)
