import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_bold_formatting(paragraph, text):
    """Simple parser to convert **bold** markdown to docx bold runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def generate_v9_report(template_path, output_path, content_dir):
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = Document(template_path)
    
    # Metadata
    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    members_list = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    members_str = "\n".join(members_list)
    year = "2026"

    # --- COVER & TITLE PAGE ---
    for p in doc.paragraphs:
        if "(Fonte: Arial ou Times 14)" in p.text:
            p.text = title.upper()
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Nome dos integrantes" in p.text:
            p.text = members_str
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "TÍTULO DO PROJETO" in p.text:
            p.text = title.upper()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "202X" in p.text:
            p.text = p.text.replace("202X", year)

    # --- CONTENT MAPPING ---
    mapping = {
        "1. INTRODUÇÃO": "introducao_desenvolvida.md",
        "2.1 OBJETIVOS": "objetivos_e_justificativa.md",
        "2.2 JUSTIFICATIVA": "objetivos_e_justificativa.md",
        "3. FUNDAMENTAÇÃO TEÓRICA": "fundamentacao_teorica.md",
        "4. METODOLOGIA": "metodologia.md",
        "5. RESULTADOS PRELIMINARES": "resultados_preliminares.md"
    }

    # Helper to find paragraph index by text match
    def find_section_index(header_text):
        for i, p in enumerate(doc.paragraphs):
            if header_text.upper() in p.text.upper():
                return i
        return -1

    # Load all MD contents
    md_contents = {}
    for filename in set(mapping.values()):
        path = os.path.join(content_dir, filename)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                md_contents[filename] = f.read()

    # Special handling for images
    # Using the latest beautified screenshots
    img_home = r"C:\Users\arpt1\.gemini\antigravity\brain\593d38a4-505d-4ccc-84b9-6da90c41da3a\permutas_baby_gallery_real_1773676799163.png"
    img_detail = r"C:\Users\arpt1\.gemini\antigravity\brain\593d38a4-505d-4ccc-84b9-6da90c41da3a\permutas_baby_product_detail_real_1773676832509.png"

    def inject_section(header_search, lines):
        idx = find_section_index(header_search)
        if idx == -1: return
        
        header_p = doc.paragraphs[idx]
        
        # Clear placeholder instruction if it's the next paragraph
        if idx + 1 < len(doc.paragraphs) and "clique aqui" in doc.paragraphs[idx+1].text.lower():
            doc.paragraphs[idx+1].text = ""

        current_base = header_p
        for line in lines:
            line = line.strip()
            if not line: continue
            
            # Skip headers from the MD itself as we want to use template structure
            if line.startswith('#'): continue
            
            # Handle list items
            is_list = line.startswith('* ') or line.startswith('- ')
            if is_list: line = line[2:]
            
            new_p = doc.add_paragraph()
            try:
                if is_list:
                    new_p.style = doc.styles['b) texto com bullets']
                else:
                    new_p.style = doc.styles['a) texto-base']
            except:
                new_p.style = doc.styles['Normal']
            
            apply_bold_formatting(new_p, line)
            
            # Detect image placeholders in MD
            if "![" in line:
                new_p.text = "" # Clear text if it's just an image line
                if "home" in line.lower():
                    if os.path.exists(img_home):
                        new_p.add_run().add_picture(img_home, width=Inches(6))
                elif "detail" in line.lower() or "produto" in line.lower():
                    if os.path.exists(img_detail):
                        new_p.add_run().add_picture(img_detail, width=Inches(6))

            current_base._element.addnext(new_p._element)
            current_base = new_p

    # Extract sections from MDs and inject
    for header, filename in mapping.items():
        if filename not in md_contents: continue
        content = md_contents[filename]
        
        # Filter content for specific subsections (like 2.1 vs 2.2)
        lines = []
        capture = False
        # If the header in mapping is specific (like 2.1), we only take that part
        if header.startswith("2."):
            prefix = header.split(' ')[0]
            for l in content.split('\n'):
                if l.startswith(f'# {prefix}') or l.startswith(f'## {prefix}'):
                    capture = True
                    continue
                elif capture and (l.startswith('# ') or l.startswith('## ')):
                    break
                if capture: lines.append(l)
        else:
            lines = content.split('\n')
            
        inject_section(header, lines)

    doc.save(output_path)
    print(f"SUCCESS: Report V9 generated at {output_path}")

if __name__ == "__main__":
    t = r'f:\projetos_opencode\PI2016Django\base_teorica\Modelo_-_Relatorio_Parcial_oficial.docx'
    o = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V10_Oficial.docx'
    c_dir = r'f:\projetos_opencode\PI2016Django\base_teorica'
    generate_v9_report(t, o, c_dir)
