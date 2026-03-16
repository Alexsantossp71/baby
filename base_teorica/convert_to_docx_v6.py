import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def add_spacer(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0 # Simple spacing for spacers

def add_capa(doc, title, members):
    # 1. UNIVESP Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO – UNIVESP")
    run.bold = True
    run.font.size = Pt(12)

    add_spacer(doc, 2)

    # 2. Members (Top)
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.size = Pt(12)

    add_spacer(doc, 10)

    # 3. Title (Center)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    add_spacer(doc, 12)

    # 4. Local/Year (Bottom)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SÃO PAULO\n2026")
    doc.add_page_break()

def add_folha_de_rosto(doc, title, members, orientador):
    # 1. Header or Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO")
    run.bold = True
    
    add_spacer(doc, 2)
    
    # 2. Authors
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)

    add_spacer(doc, 6)

    # 3. Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    add_spacer(doc, 4)

    # 4. Natureza do Trabalho (Indented to the right)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8) # Natureza text is usually on the right half
    run = p.add_run("Relatório Parcial apresentado na disciplina de Projeto Integrador em Computação I, para o curso de Tecnologia em Processos Gerenciais/Ciência de Dados da Universidade Virtual do Estado de São Paulo (UNIVESP).")
    run.font.size = Pt(10)
    
    add_spacer(doc, 2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(8)
    run = p.add_run(f"Orientador: {orientador}")
    run.font.size = Pt(10)
    run.bold = True

    add_spacer(doc, 8)

    # 5. Local/Year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SÃO PAULO\n2026")
    doc.add_page_break()

def create_relatorio_v6(md_path, docx_path):
    doc = Document()
    set_margins(doc)
    
    # Standard Body Style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    members = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    orientador = "Fernando Pinto Martinez"

    add_capa(doc, title, members)
    add_folha_de_rosto(doc, title, members, orientador)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    is_references = False
    skip_mode = False
    
    for line in lines:
        line = line.strip()
        if not line: continue
            
        # Skip header metadata that is already in prelim pages
        if line.startswith('# Relatório Parcial') or line.startswith('**Curso:') or line.startswith('**Identificação') or line.startswith('**Polo:') or line.startswith('**Orientador:'):
            continue
        
        # Skip Integrantes section
        if line.startswith('## Integrantes'):
            skip_mode = True
            continue
        if skip_mode:
            if line.startswith('## ') or line == '---':
                skip_mode = False
                if line == '---': continue
            else:
                continue
        
        if line == '---': continue

        if line.startswith('## '): # Section Title
            if "Referências" in line: is_references = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif line.startswith('### '): # Subtitle
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].strip())
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif line.startswith('- '): # List item
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
        else: # Body Paragraph
            p = doc.add_paragraph()
            if is_references:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(6)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Cm(1.25)
            
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True

    doc.save(docx_path)
    print(f"Relatório ABNT V6 (Capa + Folha de Rosto) salvo em: {docx_path}")

if __name__ == "__main__":
    md_file = r'f:\projetos_opencode\PI2016Django\base_teorica\relatorio_parcial_pi.md'
    docx_file = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V6_ABNT_Final.docx'
    create_relatorio_v6(md_file, docx_file)
