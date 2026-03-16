import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

def add_cover_page(doc, title, group_id, members, orientador):
    # UNIVESP Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDADE VIRTUAL DO ESTADO DE SÃO PAULO – UNIVESP")
    run.bold = True
    run.font.size = Pt(12)

    for _ in range(5): doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)

    for _ in range(4): doc.add_paragraph()

    # Relatório Parcial Text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO PARCIAL")
    run.bold = True
    run.font.size = Pt(12)

    for _ in range(4): doc.add_paragraph()

    # Group and Members
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Grupo: {group_id}")
    run.bold = True
    
    for member in members:
        p = doc.add_paragraph(member)
        p.paragraph_format.left_indent = Cm(0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Orientador: {orientador}")
    run.bold = True

    # Fill space to bottom
    for _ in range(12): doc.add_paragraph()

    # City/Year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SÃO PAULO\n2026")
    doc.add_page_break()

def create_abnt_relatorio_v4(md_path, docx_path):
    doc = Document()
    set_margins(doc)
    
    # Global Style Configuration
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    # Pre-defined Title and Data
    title = "Permutas Baby: Plataforma de Economia Circular para Artigos Infantis"
    group_id = "PJI110 - A2026S1N1 - Grupo 4"
    members = [
        "ALEXANDRE RAMOS DE PAULA (RA: 1700561)",
        "ANA CLARA BARBOSA NISHIMURA",
        "CINTHIA DE OLIVEIRA ARAUJO MONTEIRO",
        "SABRINA DE OLIVEIRA ESCOLASTICO GOMES",
        "ROSELI GONCALVES DE MIRANDA",
        "ANDRESSA SANTOS BANDEIRA COSTA",
        "THEYLON VIANNA SALES",
        "MARCUS VINICIUS MARTINS"
    ]
    orientador = "Fernando Pinto Martinez"

    add_cover_page(doc, title, group_id, members, orientador)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    is_references = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Ignore cover elements that are already handled
        if line.startswith('# Relatório Parcial') or line.startswith('**Curso:') or line.startswith('**Identificação') or line.startswith('**Polo:') or line.startswith('**Orientador:') or line.startswith('## Integrantes'):
            continue
        if line == '---' or (line.startswith('- ') and 'RA:' in line):
            continue

        if line.startswith('## '): # Section Title
            if "Referências" in line: is_references = True
            p = doc.add_paragraph()
            # ABNT: Separate by 1.5 space
            p.paragraph_format.space_before = Pt(12) 
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].upper()) # ABNT Title often in Uppercase
            run.bold = True
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            
        elif line.startswith('### '): # Subtitle
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[4:])
            run.bold = True # Title case or Bold for subtitles
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)

        elif line.startswith('- '): # List
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Process bold
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
        else: # Normal Paragraph
            p = doc.add_paragraph()
            if is_references:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6) # Spacing between refs
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.first_line_indent = Cm(1.25) # Standard 1.25cm
            
            # Simple Bold/Italic processing
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True

    doc.save(docx_path)
    print(f"Relatório ABNT V4 (Site-Refined) salvo em: {docx_path}")

if __name__ == "__main__":
    md_file = r'f:\projetos_opencode\PI2016Django\base_teorica\relatorio_parcial_pi.md'
    docx_file = r'f:\projetos_opencode\PI2016Django\base_teorica\Relatorio_Parcial_Permutas_Baby_V4_ABNT_Final.docx'
    create_abnt_relatorio_v4(md_file, docx_file)
