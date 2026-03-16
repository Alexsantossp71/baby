import os
from docx import Document
from docx.shared import Pt, Cm

def deep_inspect(path):
    if not os.path.exists(path):
        return

    doc = Document(path)
    print(f"--- Inspeção Profunda ABNT: {os.path.basename(path)} ---")
    
    # Check for Paragraph Indentation
    print("\n--- Recuos de Parágrafo (Primeira Linha) ---")
    indents = []
    for p in doc.paragraphs:
        if p.text.strip() and p.paragraph_format.first_line_indent:
            indents.append(p.paragraph_format.first_line_indent.cm)
    if indents:
        avg_indent = sum(indents) / len(indents)
        print(f"Recuo médio da primeira linha: {avg_indent:.2f}cm")
    else:
        print("Nenhum recuo de primeira linha detectado nos parágrafos.")

    # Check for alignment of body text
    print("\n--- Alinhamento e Espaçamento ---")
    for i, p in enumerate(doc.paragraphs[30:40]): # Check a middle section for body text
        if p.text.strip():
            print(f"P{i} Align: {p.paragraph_format.alignment}, LineSpacing: {p.paragraph_format.line_spacing}")

    # Check References section
    print("\n--- Seção de Referências (Final) ---")
    found_refs = False
    for p in doc.paragraphs[-15:]:
        if "Referências" in p.text:
            found_refs = True
        if found_refs and p.text.strip():
            print(f"Ref: {p.text[:100]}... [Align: {p.paragraph_format.alignment}]")

if __name__ == "__main__":
    example_path = r'f:\projetos_opencode\PI2016Django\base_teorica\EXEMPLO - DRP04-Projeto_Integrador-Relatório Parcial.docx'
    deep_inspect(example_path)
